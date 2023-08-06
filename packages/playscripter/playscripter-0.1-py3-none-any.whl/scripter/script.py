from scripter.scene import Act
import docx
from docx.enum import text
from docx.shared import Inches


class Script:
    """
    The script for the play.
    """
    def __init__(self, name: str, author: str):
        """
        New instance of `Script`.
        :param name: The name of the play
        :param author: The author of the play.
        """
        self.name = name
        self.author = author
        self.acts = []

    def add_act(self):
        """
        Appends a new instance of `Act` to
        `Script.acts`.
        :return: The newly created `Act`
        """
        self.acts.append(Act())
        return self.acts[-1]

    def create(self, filename: str):
        """
        Creates a new `.docx` (Word) document
        for the script.
        :param filename: The name of the file
        """
        doc = docx.Document()
        doc.add_heading("{} by {}".format(self.name, self.author), 0)
        for act in self.acts:
            doc.add_heading("Act {}".format(self.acts.index(act) + 1), 1)
            for scene in act.scenes:
                doc.add_heading("Scene {}".format(act.scenes.index(scene) + 1), 2)
                ad = doc.add_heading(level=3)
                ad.add_run(scene.description).italic = True
                for d in scene.dialogue:
                    if hasattr(d, 'character'):
                        doc.add_heading(d.character, 4)
                        s = doc.add_paragraph(d.description)
                        s.paragraph_format.left_indent = Inches(.25)
                    else:
                        sd = doc.add_heading(str(d), 4)
                        sd.alignment = text.WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_heading("THE END.", 2)

        doc.save("{}.docx".format(filename))
